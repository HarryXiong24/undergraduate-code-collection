import requests
from bs4 import BeautifulSoup
import xlwt
from docx import Document

def main():
	url="https://www.cuit.edu.cn/DepartmentList?id=2"
	kv={'user-agent':'chrome/10.0'}
	try:
		r=requests.get(url,headers=kv,timeout=20)
		r.raise_for_status()
		r.encoding=r.apparent_encoding
		html=r.text
	except:
		print("Error!")

	html=BeautifulSoup(html,"html5lib")
	ul=html.find("ul",id="DepartList")
	li=ul.find_all("li")

	xls=xlwt.Workbook()
	sheet=xls.add_sheet('成都信息工程大学教学部门',cell_overwrite_ok=True)
	sheet.write_merge(0,0,0,1,'教学部门')

	docx=Document()
	docx.add_heading('成都信息工程大学教学部门',0)
	table=docx.add_table(rows=20,cols=1)

	k=0
	for k in range(len(li)):
		string=li[k].text
		name=string[4:]
		Excel(name,sheet,k)
		Doc(name,docx,k,table)
		k+=1
		print(name)

	xls.save('成都信息工程大学教学部门.xls')
	docx.save('成都信息工程大学教学部门.doc')

def Excel(name,sheet,k):
	sheet.write(k+1,0,name)

def Doc(name,docx,k,table):
	p=docx.add_paragraph(name)
	hdr_cells=table.rows[k].cells
	hdr_cells[0].text=name




main()
